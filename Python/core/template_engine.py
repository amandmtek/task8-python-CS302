import os
from datetime import datetime
from typing import Dict, Any
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML


from docx import Document
from qrcode import make as make_qrcode
from qrcode.image.pil import PilImage

class TemplateEngine:
    
    def __init__(self):
        template_dir = os.path.join(os.path.dirname(__file__), '..', 'templates')
        self.env = Environment(loader=FileSystemLoader(template_dir))

    def _generate_qr_code(self, quest_id: int) -> str:
        url = f"https://adventurers-guild.com/quest/{quest_id}"
        img: PilImage = make_qrcode(url, image_factory=PilImage)
        
        import io
        import base64
        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        return base64.b64encode(buffer.getvalue()).decode()

    def render_html(self, template_name: str, quest_data: Dict[str, Any]) -> str:
        template = self.env.get_template(template_name)
        
        qr_code_base64 = self._generate_qr_code(quest_data.get('id', -1))
        
        context = {
            'quest': quest_data,
            'current_date': datetime.now().strftime("%d.%m.%Y"),
            'qr_code': qr_code_base64
        }
        return template.render(context)

    def export_pdf(self, template_name: str, quest_data: Dict[str, Any], output_path: str):
        html_content = self.render_html(template_name, quest_data)
        HTML(string=html_content).write_pdf(output_path)
        
    def export_docx(self, quest_data: Dict[str, Any], output_path: str):
        doc = Document()
        doc.add_heading(f"Контракт Гильдии Приключенцев #{quest_data.get('id', 'N/A')}", 0)
        
        doc.add_paragraph(f"Название: {quest_data.get('title', 'N/A')}")
        doc.add_paragraph(f"Сложность: {quest_data.get('difficulty', 'N/A')}")
        doc.add_paragraph(f"Вознаграждение: {quest_data.get('reward', 'N/A')} золотых")
        doc.add_paragraph(f"Описание:\n{quest_data.get('description', 'N/A')}")
        
        doc.save(output_path)

template_engine = TemplateEngine()
